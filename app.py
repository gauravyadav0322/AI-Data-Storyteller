"""
AI-Powered Data Storyteller (Streamlit + Ollama Vicuna)
Author: Gaurav Yadav

Goal:
Upload CSV → Automated EDA → Vicuna Insights via Ollama → Visualizations → Export Report
"""

import os
import json
import tempfile
from datetime import datetime
from typing import List, Optional

import requests
import streamlit as st
import pandas as pd
import numpy as np
import plotly.express as px
import plotly.graph_objects as go
from fpdf import FPDF
from docx import Document

# -------------------------------------------------------------------
# 1. CONFIGURATION
# -------------------------------------------------------------------
PAGE_TITLE = "AI-Powered Data Storyteller"
st.set_page_config(page_title=PAGE_TITLE, layout="wide", initial_sidebar_state="expanded")

# Ollama API settings
OLLAMA_URL = "http://localhost:11434/api/generate"
MODEL_NAME = "vicuna"   # Ensure model is available via: `ollama pull vicuna`

# -------------------------------------------------------------------
# 2. OLLAMA VICUNA HELPER FUNCTIONS
# -------------------------------------------------------------------

def llm_insights_ollama(prompt: str, model: str = MODEL_NAME) -> Optional[str]:
    """Call Vicuna via Ollama REST API and return text output."""
    try:
        payload = {"model": model, "prompt": prompt}
        response = requests.post(OLLAMA_URL, json=payload, stream=True)

        output = ""
        for line in response.iter_lines():
            if line:
                data = json.loads(line.decode("utf-8"))
                if "response" in data:
                    output += data["response"]

        return output.strip()

    except Exception as e:
        st.warning(f"Ollama request failed: {e}")
        return None


def rule_based_insights(df: pd.DataFrame) -> str:
    """Fallback insights if LLM is not available."""
    lines = []
    nrows, ncols = df.shape
    lines.append(f"- Dataset has {nrows:,} rows and {ncols} columns.")

    missing = df.isnull().sum()
    missing_cols = missing[missing > 0]
    if not missing_cols.empty:
        lines.append(f"- {len(missing_cols)} columns contain missing values.")

    num_cols = df.select_dtypes(include=[np.number]).columns.tolist()
    if num_cols:
        col = num_cols[0]
        lines.append(f"- Column `{col}`: mean={df[col].mean():.2f}, max={df[col].max():.2f}")

    return (
        "\n".join(lines)
        + "\n\nRecommendations:\n"
        + "1. Handle missing values.\n"
        + "2. Explore trends over time.\n"
        + "3. Drill into key categories."
    )


def generate_insights(df_sample: pd.DataFrame, quick_summary: str) -> str:
    """Generate structured insights from Ollama, fallback to rule-based."""
    prompt = (
        "You are a professional data analyst. Provide a comprehensive, detailed executive summary "
        "of the dataset.\n\n"
        "Structure your response into these sections:\n"
        "1. **Executive Summary**\n"
        "2. **Key Insights and Trends** (6–10 bullet points)\n"
        "3. **Actionable Recommendations** (3–5 points)\n"
        "4. **Data Quality and Limitations**\n\n"
        f"Dataset summary:\n{quick_summary}\n"
        "Ensure the summary is clear, professional, and targeted at a business audience."
    )

    text = llm_insights_ollama(prompt)
    return text if text else rule_based_insights(df_sample)

# -------------------------------------------------------------------
# 3. UTILITIES
# -------------------------------------------------------------------

@st.cache_data
def read_csv(uploaded_file):
    """Read CSV into a pandas DataFrame."""
    return pd.read_csv(uploaded_file)


def summarize_numeric(df: pd.DataFrame) -> pd.DataFrame:
    """Return numeric summary with missing counts."""
    numeric = df.select_dtypes(include=[np.number])
    if numeric.empty:
        return pd.DataFrame()

    desc = numeric.describe().T
    desc["missing"] = df[numeric.columns].isnull().sum().values
    return desc


def top_value_counts(df: pd.DataFrame, top_n: int = 10):
    """Return top N value counts for categorical columns."""
    cats = df.select_dtypes(exclude=[np.number])
    result = {}
    for col in cats.columns:
        result[col] = df[col].value_counts(dropna=False).head(top_n)
    return result


def create_pdf(title: str, text: str, images: List[str], out_path: str):
    """Generate a PDF report with text and images."""
    pdf = FPDF(unit="mm", format="A4")
    pdf.add_page()
    pdf.set_font("Arial", "B", 16)
    pdf.cell(0, 12, title, ln=True, align="C")
    pdf.ln(4)

    pdf.set_font("Arial", size=11)
    for line in text.split("\n"):
        pdf.multi_cell(0, 6, line)

    for img in images:
        try:
            pdf.add_page()
            pdf.image(img, x=10, y=30, w=180)
        except Exception:
            pass

    pdf.output(out_path)


def create_docx(title: str, text: str, images: List[str], out_path: str):
    """Generate a DOCX report with text and images."""
    doc = Document()
    doc.add_heading(title, level=1)

    for line in text.split("\n"):
        doc.add_paragraph(line)

    for img in images:
        try:
            doc.add_picture(img)
        except Exception:
            pass

    doc.save(out_path)

# -------------------------------------------------------------------
# 4. STREAMLIT APPLICATION
# -------------------------------------------------------------------

st.title(PAGE_TITLE)
st.caption("Upload a CSV → Automated EDA → Vicuna Insights (Ollama) → Visualizations → Report Export")

# --- File Upload ---
uploaded = st.file_uploader("Upload CSV", type="csv")
if not uploaded:
    st.stop()

df = read_csv(uploaded)
if df.empty:
    st.error("CSV is empty!")
    st.stop()

# --- Dataset Overview ---
st.subheader("Dataset Overview")
st.write(f"Rows: {df.shape[0]}, Columns: {df.shape[1]}")
st.dataframe(df.head())

# --- Key Metrics ---
st.subheader("Key Metrics")
k1, k2, k3 = st.columns(3)
with k1:
    st.metric("Rows", f"{df.shape[0]:,}")
with k2:
    st.metric("Columns", f"{df.shape[1]}")
with k3:
    st.metric("Missing Values", f"{df.isnull().sum().sum():,}")

# --- EDA ---
st.subheader("EDA")
numeric_summary = summarize_numeric(df)
if not numeric_summary.empty:
    st.dataframe(numeric_summary)
else:
    st.write("No numeric columns found.")

top_vals = top_value_counts(df, top_n=5)
for col, series in top_vals.items():
    st.write(f"Top values for **{col}**")
    st.dataframe(series)

# --- Visualizations ---
st.subheader("Visualizations")
num_cols = df.select_dtypes(include=[np.number]).columns.tolist()
cat_cols = df.select_dtypes(exclude=[np.number]).columns.tolist()

if num_cols:
    fig = px.histogram(df, x=num_cols[0], nbins=40, title=f"Distribution of {num_cols[0]}")
    st.plotly_chart(fig, use_container_width=True)

if cat_cols:
    fig = px.bar(df[cat_cols[0]].value_counts().head(10), title=f"Top {cat_cols[0]} values")
    st.plotly_chart(fig, use_container_width=True)

if len(num_cols) >= 2:
    corr = df[num_cols].corr()
    fig = go.Figure(data=go.Heatmap(
        z=corr.values,
        x=corr.columns,
        y=corr.index,
        colorscale="RdBu",
        zmin=-1,
        zmax=1
    ))
    fig.update_layout(title="Correlation Heatmap")
    st.plotly_chart(fig, use_container_width=True)

# --- AI Insights ---
st.subheader("AI Insights (Vicuna via Ollama)")
summary_lines = [f"Rows: {df.shape[0]}, Columns: {df.shape[1]}"]

for col in num_cols[:3]:
    summary_lines.append(f"{col}: mean={df[col].mean():.2f}, max={df[col].max():.2f}")
for col in cat_cols[:2]:
    top_vals = df[col].value_counts().head(3).to_dict()
    summary_lines.append(f"{col}: {top_vals}")

quick_summary = "\n".join(summary_lines)
insights = generate_insights(df, quick_summary)
st.write(insights)

# --- Export Report ---
st.subheader("Export Report")
tmpdir = tempfile.mkdtemp()
title = f"Executive Summary - {datetime.now().strftime('%Y-%m-%d')}"
image_paths = []

if st.button("Generate PDF"):
    out_pdf = os.path.join(tmpdir, "summary.pdf")
    create_pdf(title, insights, image_paths, out_pdf)
    with open(out_pdf, "rb") as f:
        st.download_button("Download PDF", f, "summary.pdf")

if st.button("Generate DOCX"):
    out_doc = os.path.join(tmpdir, "summary.docx")
    create_docx(title, insights, image_paths, out_doc)
    with open(out_doc, "rb") as f:
        st.download_button("Download DOCX", f, "summary.docx")
